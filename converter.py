"""文件格式转换核心模块"""
import os
import io
import zipfile
from pathlib import Path
from typing import List, Tuple, Optional
import tempfile
import shutil

# PDF
import fitz  # PyMuPDF
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Image
from PIL import Image

OUTPUT_DIR = Path(__file__).parent / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


def _ensure_tesseract():
    """确保 Tesseract 路径已配置（每次调用前执行）"""
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        candidates = [
            r"D:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for path in candidates:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                return True
        return False

_init_ok = _ensure_tesseract()


def _clean_filename(name: str) -> str:
    """清理文件名中的非法字符"""
    for ch in ['\\', '/', ':', '*', '?', '"', '<', '>', '|']:
        name = name.replace(ch, '_')
    return name


def _get_unique_path(filename: str) -> Path:
    """生成唯一输出路径，避免覆盖"""
    base = OUTPUT_DIR / filename
    if not base.exists():
        return base
    stem = base.stem
    suffix = base.suffix
    counter = 1
    while True:
        new_name = f"{stem}_{counter}{suffix}"
        new_path = OUTPUT_DIR / new_name
        if not new_path.exists():
            return new_path
        counter += 1


# ==================== PDF → 图像 ====================

def pdf_to_images(
    pdf_bytes: bytes,
    output_format: str = "png",
    dpi: int = 200,
    page_range: str = "all"
) -> List[Path]:
    """
    将PDF页面转换为图像
    page_range: "all" 或 "1,3,5" 或 "1-5"
    返回生成的文件路径列表
    """
    fmt = output_format.lower()
    ext_map = {"png": "png", "jpg": "jpg", "jpeg": "jpg", "webp": "webp",
               "tiff": "tiff", "tif": "tiff", "bmp": "bmp"}
    ext = ext_map.get(fmt, "png")
    
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    output_paths = []
    
    # 解析页码范围
    if page_range.strip().lower() == "all":
        pages = list(range(len(doc)))
    else:
        pages = _parse_page_range(page_range, len(doc))
    
    for page_num in pages:
        page = doc[page_num]
        # 使用矩阵提高分辨率
        mat = fitz.Matrix(dpi/72, dpi/72)
        pix = page.get_pixmap(matrix=mat)
        
        # 转换为PIL图像以支持更多格式
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        if ext == "jpg":
            img = img.convert("RGB")
        
        filename = f"page_{page_num + 1}.{ext}"
        out_path = _get_unique_path(filename)
        img.save(out_path, quality=95 if ext == "jpg" else None)
        output_paths.append(out_path)
    
    doc.close()
    return output_paths


def pdf_to_images_zip(
    pdf_bytes: bytes,
    output_format: str = "png",
    dpi: int = 200,
    page_range: str = "all"
) -> Path:
    """将PDF页面转为图像并打包为ZIP"""
    paths = pdf_to_images(pdf_bytes, output_format, dpi, page_range)
    if not paths:
        return None
    
    zip_path = _get_unique_path("converted_images.zip")
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for p in paths:
            zf.write(p, p.name)
    
    # 清理单独的图片文件，只保留zip
    for p in paths:
        p.unlink()
    
    return zip_path


# ==================== 图像 → PDF ====================

def images_to_pdf(
    image_files: List[bytes],
    image_names: List[str],
    output_mode: str = "merge"
) -> Path:
    """
    将图像转为PDF
    output_mode: "merge" 合并为单PDF, "separate" 各自独立
    """
    if output_mode == "separate":
        # 各自独立输出，打包zip
        pdf_paths = []
        for img_bytes, name in zip(image_files, image_names):
            img = Image.open(io.BytesIO(img_bytes))
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            
            stem = Path(name).stem
            out_path = _get_unique_path(f"{stem}.pdf")
            img.save(out_path, "PDF", resolution=100.0)
            pdf_paths.append(out_path)
        
        if len(pdf_paths) == 1:
            return pdf_paths[0]
        
        zip_path = _get_unique_path("converted_pdfs.zip")
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for p in pdf_paths:
                zf.write(p, p.name)
                p.unlink()
        return zip_path
    else:
        # 合并为单个PDF
        doc = fitz.open()
        for img_bytes, name in zip(image_files, image_names):
            img = Image.open(io.BytesIO(img_bytes))
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            
            # 临时保存为PNG再插入PDF（保持质量）
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                img.save(tmp.name, "PNG")
                tmp_path = tmp.name
            
            rect = fitz.Rect(0, 0, img.width, img.height)
            page = doc.new_page(width=img.width, height=img.height)
            page.insert_image(rect, filename=tmp_path)
            os.unlink(tmp_path)
        
        out_path = _get_unique_path("merged_images.pdf")
        doc.save(out_path)
        doc.close()
        return out_path


# ==================== PDF → Word ====================

def _extract_text_rebuild(pdf_bytes: bytes) -> tuple[str, Path]:
    """
    用 PyMuPDF 提取文本并用 python-docx 重建文档。
    返回 (提取到的文本总内容, 输出文件路径)。
    如果提取到的文本太少，返回 ("", None) 让调用方 fallback 到图片模式。
    """
    doc_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_text = ""
    all_blocks = []
    
    for page in doc_pdf:
        # 按块提取文本，保留位置信息以便判断标题/正文
        blocks = page.get_text("dict").get("blocks", [])
        page_blocks = []
        for b in blocks:
            if b.get("type") != 0:  # 跳过图片块
                continue
            for line in b.get("lines", []):
                line_text = ""
                max_size = 0
                for span in line.get("spans", []):
                    line_text += span.get("text", "")
                    max_size = max(max_size, span.get("size", 11))
                if line_text.strip():
                    page_blocks.append((line_text.strip(), max_size, b["bbox"][1]))
                    total_text += line_text + "\n"
        all_blocks.append(page_blocks)
    doc_pdf.close()
    
    # 阈值：如果总字符数太少，认为是扫描件，回退到图片模式
    if len(total_text.strip()) < 100:
        return "", None
    
    # 重建 Word 文档
    docx = Document()
    
    # 统计字体大小分布，找出标题和正文的分界线
    all_sizes = []
    for page_blocks in all_blocks:
        for _, size, _ in page_blocks:
            all_sizes.append(size)
    
    if all_sizes:
        avg_size = sum(all_sizes) / len(all_sizes)
        max_size = max(all_sizes)
        # 大于平均 + 20% 且接近最大值的认为是标题
        title_threshold = max(avg_size * 1.2, max_size * 0.75)
    else:
        title_threshold = 14
    
    for page_idx, page_blocks in enumerate(all_blocks):
        if page_idx > 0:
            docx.add_page_break()
        
        for text, size, y in page_blocks:
            # 启发式：大字体 + 短文本 = 标题
            is_title = size >= title_threshold and len(text) < 80
            
            if is_title:
                # 根据字体大小判断标题级别
                if size >= title_threshold + 4:
                    level = 1
                elif size >= title_threshold + 2:
                    level = 2
                else:
                    level = 3
                heading = docx.add_heading(text, level=level)
            else:
                p = docx.add_paragraph(text)
                # 设置正文字体
                for run in p.runs:
                    run.font.size = Pt(min(max(size, 10), 12))
    
    out_path = _get_unique_path("converted.docx")
    docx.save(out_path)
    return total_text, out_path


def pdf_to_word(pdf_bytes: bytes) -> tuple[Path, str]:
    """
    PDF 转换为 Word 文档。
    优先提取文字重建可编辑 docx；若文本太少（扫描件）则 fallback 到 pdf2docx 图片模式。
    返回 (输出路径, 模式说明)。
    """
    # 先尝试文本提取重建
    text_content, text_path = _extract_text_rebuild(pdf_bytes)
    
    if text_path and len(text_content.strip()) >= 100:
        return text_path, "文字提取模式（可编辑）"
    
    # Fallback: 扫描件/图片 PDF，用 pdf2docx 的图片模式
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
        tmp_pdf.write(pdf_bytes)
        tmp_pdf_path = tmp_pdf.name
    
    out_path = _get_unique_path("converted.docx")
    
    try:
        cv = Converter(tmp_pdf_path)
        cv.convert(str(out_path), start=0, end=None)
        cv.close()
    finally:
        os.unlink(tmp_pdf_path)
    
    return out_path, "图片模式（扫描件，无法编辑文字）"


# ==================== 图像格式互转 ====================

def convert_image(
    image_bytes: bytes,
    input_name: str,
    output_format: str
) -> Path:
    """图像格式互转"""
    fmt_map = {
        "png": "PNG", "jpg": "JPEG", "jpeg": "JPEG",
        "webp": "WEBP", "bmp": "BMP", "tiff": "TIFF", "gif": "GIF"
    }
    pil_fmt = fmt_map.get(output_format.lower(), "PNG")
    ext = output_format.lower()
    
    img = Image.open(io.BytesIO(image_bytes))
    
    # 处理透明通道
    if pil_fmt == "JPEG" and img.mode in ("RGBA", "P"):
        # 创建白色背景
        background = Image.new("RGB", img.size, (255, 255, 255))
        if img.mode == "P":
            img = img.convert("RGBA")
        background.paste(img, mask=img.split()[-1])
        img = background
    
    stem = Path(input_name).stem
    out_path = _get_unique_path(f"{stem}.{ext}")
    
    save_kwargs = {}
    if pil_fmt == "JPEG":
        save_kwargs["quality"] = 95
    elif pil_fmt == "WEBP":
        save_kwargs["quality"] = 90
    
    img.save(out_path, pil_fmt, **save_kwargs)
    return out_path


# ==================== 图像 → Word ====================

def image_to_word(
    image_bytes: bytes,
    image_name: str,
    mode: str = "embed"
) -> tuple[Path, str]:
    """
    图像转 Word
    mode: "embed"  图片嵌入（保留原图，不可编辑文字）
          "ocr"    OCR文字提取（可编辑，需安装Tesseract）
          "hybrid" 图片+文字混合（第1页原图，第2页可编辑文字）
    返回 (输出路径, 模式说明)
    """
    if mode == "ocr":
        return _image_to_word_ocr(image_bytes, image_name)
    elif mode == "hybrid":
        return _image_to_word_hybrid(image_bytes, image_name)
    else:
        return _image_to_word_embed(image_bytes, image_name), "图片嵌入模式（保留原图）"


def _image_to_word_embed(image_bytes: bytes, image_name: str) -> Path:
    """将图片直接嵌入 Word 文档"""
    doc = Document()
    
    # python-docx 的 add_picture 需要文件路径，先写临时文件
    suffix = Path(image_name).suffix or ".png"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(image_bytes)
        tmp_path = tmp.name
    
    try:
        # 计算合适的显示宽度（最大 6 英寸）
        img = Image.open(io.BytesIO(image_bytes))
        width_inch = min(img.width / 150, 6.0)  # 按 150 DPI 估算，不超过 6 英寸
        doc.add_picture(tmp_path, width=Inches(width_inch))
    finally:
        os.unlink(tmp_path)
    
    stem = Path(image_name).stem
    out_path = _get_unique_path(f"{stem}.docx")
    doc.save(out_path)
    return out_path


def _try_rapidocr(img: Image.Image) -> list:
    """尝试用 RapidOCR 识别，返回统一格式的 words 列表（纯 Python，无需 Tesseract）"""
    try:
        from rapidocr_onnxruntime import RapidOCR
        import numpy as np
        engine = RapidOCR()
        result, _ = engine(np.array(img))
        if not result:
            return []
        words = []
        for item in result:
            bbox, text, conf = item[0], item[1], item[2]
            if not text or not text.strip():
                continue
            xs = [p[0] for p in bbox]
            ys = [p[1] for p in bbox]
            words.append({
                'text': text.strip(),
                'left': min(xs),
                'top': min(ys),
                'width': max(xs) - min(xs),
                'height': max(ys) - min(ys),
                'conf': conf,
            })
        return words
    except Exception:
        return []


def _try_pytesseract(img: Image.Image, lang: str) -> list:
    """尝试用 pytesseract 识别，返回统一格式的 words 列表（需要本地 Tesseract）"""
    if not _ensure_tesseract():
        return []
    try:
        import pytesseract
        data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT)
        words = []
        for i in range(len(data['text'])):
            conf = int(data['conf'][i])
            if conf < 30:
                continue
            text = data['text'][i].strip()
            if not text:
                continue
            words.append({
                'text': text,
                'left': data['left'][i],
                'top': data['top'][i],
                'width': data['width'][i],
                'height': data['height'][i],
                'conf': conf,
            })
        return words
    except Exception:
        return []


def _build_docx_from_words(words: list, img_width: int) -> Document:
    """根据 words 列表构建 Word 文档（共用排版分析逻辑）"""
    if not words:
        return None
    
    avg_height_global = sum(w['height'] for w in words) / len(words)
    title_threshold = max(avg_height_global * 1.3, 20)
    
    # 按行分组（y 坐标接近的归为一行）
    words_sorted = sorted(words, key=lambda w: w['top'])
    lines = []
    current_line = [words_sorted[0]]
    
    for w in words_sorted[1:]:
        avg_height = sum(x['height'] for x in current_line) / len(current_line)
        if abs(w['top'] - current_line[0]['top']) < avg_height * 0.8:
            current_line.append(w)
        else:
            current_line.sort(key=lambda x: x['left'])
            lines.append(current_line)
            current_line = [w]
    current_line.sort(key=lambda x: x['left'])
    lines.append(current_line)
    
    doc = Document()
    for line_words in lines:
        line_text = ''.join(w['text'] for w in line_words)
        avg_height = sum(w['height'] for w in line_words) / len(line_words)
        leftmost = line_words[0]['left']
        rightmost = line_words[-1]['left'] + line_words[-1]['width']
        line_center = (leftmost + rightmost) / 2
        
        is_center = abs(line_center - img_width / 2) < img_width * 0.12
        is_title = avg_height >= title_threshold and len(line_text) < 60
        
        if is_title:
            level = 1 if avg_height >= title_threshold * 1.2 else 2
            p = doc.add_heading(line_text, level=level)
        else:
            p = doc.add_paragraph(line_text)
        
        if is_center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


def _ocr_with_layout(image_bytes: bytes, image_name: str, lang: str = "chi_sim+eng") -> tuple[Document, str]:
    """
    OCR + 版面分析，返回 (Document对象, 模式说明)。
    优先 RapidOCR（云端友好），回退 pytesseract（本地）。
    """
    img = Image.open(io.BytesIO(image_bytes))
    img_width = img.width
    
    # 先尝试 RapidOCR（纯 Python ONNX，无需 Tesseract，适合 Streamlit Cloud）
    words = _try_rapidocr(img)
    engine_name = "RapidOCR"
    
    # 失败则尝试 pytesseract（需要本地安装 Tesseract）
    if not words:
        words = _try_pytesseract(img, lang)
        engine_name = "Tesseract"
    
    if not words:
        return None, "OCR未识别到文字"
    
    doc = _build_docx_from_words(words, img_width)
    if doc is None:
        return None, "OCR未识别到文字"
    
    return doc, f"{engine_name}版面分析模式（可编辑）"


def _image_to_word_ocr(image_bytes: bytes, image_name: str, lang: str = "chi_sim+eng") -> tuple[Path, str]:
    """OCR 提取文字生成可编辑 Word（带版面分析）"""
    doc, info = _ocr_with_layout(image_bytes, image_name, lang)
    
    if doc is None:
        path = _image_to_word_embed(image_bytes, image_name)
        return path, f"图片嵌入模式（{info}）"
    
    stem = Path(image_name).stem
    out_path = _get_unique_path(f"{stem}_ocr.docx")
    doc.save(out_path)
    return out_path, info


def _image_to_word_hybrid(image_bytes: bytes, image_name: str, lang: str = "chi_sim+eng") -> tuple[Path, str]:
    """
    混合模式：第1页嵌入原图（保留格式+图片），第2页附上OCR可编辑文字（带版面分析）
    """
    doc, info = _ocr_with_layout(image_bytes, image_name, lang)
    
    out_doc = Document()
    
    # ===== 第1页：嵌入原图 =====
    suffix = Path(image_name).suffix or ".png"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(image_bytes)
        tmp_path = tmp.name
    
    try:
        img_obj = Image.open(io.BytesIO(image_bytes))
        width_inch = min(img_obj.width / 150, 6.0)
        out_doc.add_picture(tmp_path, width=Inches(width_inch))
    finally:
        os.unlink(tmp_path)
    
    # ===== 第2页：OCR 文字 =====
    if doc is not None:
        out_doc.add_page_break()
        out_doc.add_heading("【OCR 文字提取结果 - 可编辑】", level=1)
        
        # 复制段落
        for element in doc.element.body:
            out_doc.element.body.append(element)
    
    stem = Path(image_name).stem
    out_path = _get_unique_path(f"{stem}_hybrid.docx")
    out_doc.save(out_path)
    return out_path, f"图片+文字混合模式（{info}）"


# ==================== 图像合并 ====================

def merge_images_vertical(
    image_files: List[bytes],
    image_names: List[str],
    spacing: int = 0,
    bg_color: tuple = (255, 255, 255),
    align: str = "center",
    max_width: int = None
) -> Path:
    """
    将多张图片垂直拼接为一张长图
    spacing: 图片间距（像素）
    bg_color: 背景色 (R, G, B)
    align: left/center/right
    max_width: 最大宽度，超出则等比缩放
    """
    images = []
    for img_bytes in image_files:
        img = Image.open(io.BytesIO(img_bytes))
        if img.mode in ("RGBA", "P"):
            # 转成 RGB，透明部分用背景色填充
            background = Image.new("RGB", img.size, bg_color)
            if img.mode == "P":
                img = img.convert("RGBA")
            background.paste(img, mask=img.split()[-1] if img.mode == "RGBA" else None)
            img = background
        elif img.mode != "RGB":
            img = img.convert("RGB")
        images.append(img)
    
    # 确定目标宽度
    target_width = max(img.width for img in images)
    if max_width and target_width > max_width:
        target_width = max_width
    
    # 缩放所有图片到统一宽度（保持比例）
    scaled = []
    for img in images:
        if img.width != target_width:
            ratio = target_width / img.width
            new_height = int(img.height * ratio)
            img = img.resize((target_width, new_height), Image.LANCZOS)
        scaled.append(img)
    
    # 计算总高度
    total_height = sum(img.height for img in scaled) + spacing * (len(scaled) - 1)
    
    # 创建画布
    result = Image.new("RGB", (target_width, total_height), bg_color)
    
    # 逐张粘贴
    y = 0
    for img in scaled:
        x = 0
        if align == "center":
            x = (target_width - img.width) // 2
        elif align == "right":
            x = target_width - img.width
        result.paste(img, (x, y))
        y += img.height + spacing
    
    out_path = _get_unique_path("merged_long_image.png")
    result.save(out_path, "PNG", quality=95)
    return out_path


# ==================== PDF 合并 ====================

def merge_pdfs(pdf_files: List[bytes], pdf_names: List[str]) -> Path:
    """合并多个PDF为一个"""
    doc = fitz.open()
    tmp_files = []
    
    try:
        for pdf_bytes in pdf_files:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(pdf_bytes)
                tmp_files.append(tmp.name)
            
            src = fitz.open(tmp.name)
            doc.insert_pdf(src)
            src.close()
        
        out_path = _get_unique_path("merged.pdf")
        doc.save(out_path)
        return out_path
    finally:
        doc.close()
        for f in tmp_files:
            if os.path.exists(f):
                os.unlink(f)


# ==================== PDF 拆分 ====================

def split_pdf(pdf_bytes: bytes, split_mode: str, ranges: str = "") -> Path:
    """
    拆分PDF
    split_mode: "single" 每页单独, "range" 按范围
    ranges: 如 "1-3,5,7-9"
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total = len(doc)
    output_files = []
    
    try:
        if split_mode == "single":
            for i in range(total):
                new_doc = fitz.open()
                new_doc.insert_pdf(doc, from_page=i, to_page=i)
                out_path = _get_unique_path(f"page_{i+1}.pdf")
                new_doc.save(out_path)
                new_doc.close()
                output_files.append(out_path)
        else:
            ranges_list = _parse_split_ranges(ranges, total)
            for idx, (start, end) in enumerate(ranges_list):
                new_doc = fitz.open()
                new_doc.insert_pdf(doc, from_page=start-1, to_page=end-1)
                out_path = _get_unique_path(f"part_{idx+1}.pdf")
                new_doc.save(out_path)
                new_doc.close()
                output_files.append(out_path)
        
        if len(output_files) == 1:
            return output_files[0]
        
        # 打包为zip
        zip_path = _get_unique_path("split_pdfs.zip")
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for p in output_files:
                zf.write(p, p.name)
                p.unlink()
        return zip_path
    finally:
        doc.close()


# ==================== PDF 提取文本 ====================

def extract_pdf_text(pdf_bytes: bytes) -> Path:
    """提取PDF文本保存为txt"""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    texts = []
    for page in doc:
        texts.append(f"--- 第 {page.number + 1} 页 ---\n")
        texts.append(page.get_text())
        texts.append("\n")
    doc.close()
    
    out_path = _get_unique_path("extracted_text.txt")
    out_path.write_text("\n".join(texts), encoding="utf-8")
    return out_path


# ==================== 工具函数 ====================

def _parse_page_range(page_range: str, total_pages: int) -> List[int]:
    """解析页码范围字符串，返回0-based页码列表"""
    pages = set()
    parts = page_range.split(",")
    for part in parts:
        part = part.strip()
        if "-" in part:
            start, end = part.split("-")
            start = max(1, int(start.strip()))
            end = min(total_pages, int(end.strip()))
            pages.update(range(start - 1, end))
        else:
            p = int(part.strip())
            if 1 <= p <= total_pages:
                pages.add(p - 1)
    return sorted(list(pages))


def _parse_split_ranges(ranges: str, total_pages: int) -> List[Tuple[int, int]]:
    """解析拆分范围，返回 (start, end) 列表，1-based"""
    result = []
    parts = ranges.split(",")
    for part in parts:
        part = part.strip()
        if "-" in part:
            start, end = part.split("-")
            start = max(1, int(start.strip()))
            end = min(total_pages, int(end.strip()))
            result.append((start, end))
        else:
            p = int(part.strip())
            if 1 <= p <= total_pages:
                result.append((p, p))
    return result


def cleanup_old_files(max_age_hours: int = 24):
    """清理超过指定时间的旧文件"""
    import time
    now = time.time()
    max_age = max_age_hours * 3600
    for f in OUTPUT_DIR.iterdir():
        if f.is_file() and (now - f.stat().st_mtime) > max_age:
            f.unlink()
